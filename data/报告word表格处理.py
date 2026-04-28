import sys
import os
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

"""
Word表格一键优化工具 - 批量优化Word文档中的表格样式。

功能：
    - 批量处理Word文档中的表格
    - 自动应用专业表格样式（边框、字体、对齐等）

使用方式：
    from data.报告word表格处理 import process_word_tables
    process_word_tables(file_list, output_dir)

依赖：python-docx
"""

import os
import re
from docx import Document
from docx.shared import Pt, Twips
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from modules.logger import log_output


def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    try:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        for edge in ("top", "bottom", "left", "right"):
            edge_data = kwargs.get(edge)
            
            # 先移除已有的该方向边框
            # 使用正确的方式查找和移除边框元素
            border_tag = qn(f"w:{edge}")
            for child in list(tcPr):
                if child.tag == border_tag:
                    tcPr.remove(child)
            
            if edge_data:
                tag = f"w:{edge}"
                try:
                    element = OxmlElement(tag)
                    tcPr.append(element)
                    for key, value in edge_data.items():
                        element.set(qn(f"w:{key}"), str(value))
                except Exception:
                    pass
    except Exception:
        pass


def find_header_end_row(table):
    """找到表头结束行（以首行纵向合并最大的单元格为准）"""
    n_rows = len(table.rows)
    
    if n_rows == 0:
        return 1
    
    max_span = 1
    first_row = table.rows[0]
    
    # 检查第一行的每个单元格是否有垂直合并标记
    for col_idx, cell in enumerate(first_row.cells):
        v_merge = cell._tc.get_or_add_tcPr().find(qn("w:vMerge"))
        
        if v_merge is not None:
            v_merge_val = v_merge.get(qn("w:val"))
            
            # 如果单元格标记为合并开始
            if v_merge_val is None or v_merge_val == "" or v_merge_val == "restart":
                current_span = 1
                
                # 向下检查：如果下一行相同位置也是restart标记，说明表头跨越多行
                if n_rows > 1:
                    next_row = table.rows[1]
                    if col_idx < len(next_row.cells):
                        next_cell = next_row.cells[col_idx]
                        next_v_merge = next_cell._tc.get_or_add_tcPr().find(qn("w:vMerge"))
                        
                        if next_v_merge is not None:
                            next_v_merge_val = next_v_merge.get(qn("w:val"))
                            # 如果下一行也是restart或空值，说明表头至少跨越两行
                            if next_v_merge_val is None or next_v_merge_val == "" or next_v_merge_val == "restart":
                                current_span = 2
                                
                                # 继续检查第三行
                                if n_rows > 2:
                                    third_row = table.rows[2]
                                    if col_idx < len(third_row.cells):
                                        third_cell = third_row.cells[col_idx]
                                        third_v_merge = third_cell._tc.get_or_add_tcPr().find(qn("w:vMerge"))
                                        if third_v_merge is not None:
                                            third_v_merge_val = third_v_merge.get(qn("w:val"))
                                            if third_v_merge_val is None or third_v_merge_val == "" or third_v_merge_val == "restart":
                                                current_span = 3
                
                if current_span > max_span:
                    max_span = current_span
    
    return max_span


def apply_table_style(table):
    """将专业表格样式应用到给定的表格对象上"""
    n_rows = len(table.rows)
    n_cols = len(table.columns)

    # 找到表头结束行（考虑合并单元格）
    header_end_row = find_header_end_row(table)

    # 清除所有边框
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            set_cell_border(cell, top={}, bottom={}, left={}, right={})

    # 设置内边框（虚线0.5磅）- 只设置左边框
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            if j > 0:
                set_cell_border(cell, left={"sz": "8", "val": "dotted"})

    # 设置标题行底端边框（实线0.5磅）- 使用表头结束行，确保整行都是实线
    if n_rows > 0 and header_end_row > 0 and header_end_row <= n_rows:
        header_row = table.rows[header_end_row - 1]
        for cell in header_row.cells:
            set_cell_border(cell, bottom={"sz": "8", "val": "single"})
        # 设置下一行的顶部边框
        if header_end_row < n_rows:
            next_row = table.rows[header_end_row]
            for cell in next_row.cells:
                set_cell_border(cell, top={"sz": "8", "val": "single"})

    # 最后处理：为纵向合并单元格设置上方1磅实线和下方0.5磅实线
    # 遍历所有行，找到所有纵向合并单元格
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            v_merge = cell._tc.get_or_add_tcPr().find(qn("w:vMerge"))
            if v_merge is not None:
                v_merge_val = v_merge.get(qn("w:val"))
                # 如果是合并标记（开始或继续）
                if v_merge_val is None or v_merge_val == "" or v_merge_val == "restart" or v_merge_val == "continue":
                    # 设置上方1磅实线（只在第一行设置）
                    if i == 0:
                        set_cell_border(cell, top={"sz": "16", "val": "single"})
                    
                    # 设置下方0.5磅实线
                    set_cell_border(cell, bottom={"sz": "8", "val": "single"})
                    
                    # 设置下一行对应单元格的顶部边框（确保显示实线）
                    if i + 1 < n_rows:
                        next_row = table.rows[i + 1]
                        if j < len(next_row.cells):
                            next_cell = next_row.cells[j]
                            set_cell_border(next_cell, top={"sz": "8", "val": "single"})

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    text = run.text
                    if text and not any('一' <= c <= '鿿' for c in text):
                        run.font.name = 'Times New Roman'
                    else:
                        run.font.name = 'Arial Narrow'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
                    run.font.size = Pt(10.5)
                paragraph.paragraph_format.line_spacing = Pt(20)
                paragraph.paragraph_format.line_spacing_rule = 3

    if n_rows > 0:
        # 表头所有行都居中（从第一行到header_end_row）
        for i in range(min(header_end_row, n_rows)):
            for cell in table.rows[i].cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.bold = False

    for row in table.rows[1:]:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for row in table.rows[1:]:
        for cell in row.cells:
            text = cell.text.strip()
            if text == "合计":
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                continue

            cleaned = re.sub(r',', '', text)
            if cleaned.replace('.', '', 1).isdigit():
                try:
                    num = float(cleaned)
                    parts = text.split('.')
                    decimals = len(parts[1]) if len(parts) > 1 else 0
                    format_str = f"{{:,.{decimals}f}}" if decimals > 0 else "{:,.0f}"
                    formatted = format_str.format(num)
                    if formatted != text:
                        cell.text = formatted
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(10.5)
                except:
                    pass


def process_word_tables(file_list, output_dir):
    """
    批量处理Word文档中的表格

    Args:
        file_list: 文件路径列表
        output_dir: 输出目录

    Returns:
        tuple: (成功数, 失败数)
    """
    success_count = 0
    fail_count = 0
    total = len(file_list)

    log_output(f"开始批量处理：共 {total} 个文件")
    log_output(f"输出目录：{output_dir}")

    for idx, src_path in enumerate(file_list, 1):
        base = os.path.basename(src_path)
        name, ext = os.path.splitext(base)

        log_output(f"[{idx}/{total}] 正在处理: {base}")

        try:
            doc = Document(src_path)
            for table in doc.tables:
                apply_table_style(table)

            os.makedirs(output_dir, exist_ok=True)
            dst_path = os.path.join(output_dir, f"{name}_优化后{ext}")
            doc.save(dst_path)

            success_count += 1
            log_output(f"✓ 成功: {base} -> {os.path.basename(dst_path)}")

        except Exception as e:
            fail_count += 1
            import traceback
            log_output(f"✗ 失败: {base}")
            log_output(f"错误: {str(e)}")
            log_output(f"详细: {traceback.format_exc()}")

    log_output(f"处理完成：成功 {success_count} / 总计 {total}")

    return success_count, fail_count


if __name__ == "__main__":
    import tkinter as tk
    from tkinter import filedialog, messagebox
    import threading

    root = tk.Tk()
    root.withdraw()

    files = filedialog.askopenfilenames(
        title="选择Word文档",
        filetypes=[("Word文档", "*.docx")]
    )

    if not files:
        exit(0)

    output_dir = filedialog.askdirectory(title="选择输出目录")
    if not output_dir:
        exit(0)

    def run():
        success, fail = process_word_tables(list(files), output_dir)
        root.after(0, lambda: messagebox.showinfo("完成", f"成功 {success}，失败 {fail}"))
        root.after(0, root.destroy)

    thread = threading.Thread(target=run)
    thread.start()
    root.mainloop()